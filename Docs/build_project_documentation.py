# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
GENERATED_DATE = "June 26, 2026"
BRANCH = "Indev-3-Feature-CloneUtilityRedesign"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(88, 88, 88)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F6F8FA"
BORDER = "DADCE0"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def configure_document(doc: Document, title: str):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 8),
        ("Heading 2", 13, BLUE, 14, 6),
        ("Heading 3", 11.5, DARK_BLUE, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.text = title
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header_p.runs[0], size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = f"Generated {GENERATED_DATE} - {BRANCH}"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer_p.runs[0], size=8.5, color=MUTED)


def add_title_page(doc: Document, title: str, subtitle: str, version_label: str, scope: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=24, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    set_run_font(r, size=13.5, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    rows = [
        ("Version", version_label),
        ("Generated", GENERATED_DATE),
        ("Branch", BRANCH),
        ("Scope", scope),
    ]
    for i, (label, value) in enumerate(rows):
        cells = meta.rows[i].cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_cell_width(cells[0], 1900)
        set_cell_width(cells[1], 7460)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, color=INK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(8)
    r = note.add_run(
        "This is a living technical handoff. It describes how the current project operates, why major design decisions were made, and where future work should continue."
    )
    set_run_font(r, italic=True, color=MUTED)
    add_border_bottom(note, color="DADCE0", size="6")
    doc.add_page_break()


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_code_board(doc, title, board):
    add_h3(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 9360)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_cell_shading(cell, CODE_FILL)
    remove_cell_borders(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for i, line in enumerate(board.strip("\n").split("\n")):
        if i > 0:
            p.add_run("\n")
        r = p.add_run(line.rstrip())
        set_run_font(r, name="Consolas", size=8.5, color=RGBColor(25, 25, 25))
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_margins(hdr[i])
        for run in hdr[i].paragraphs[0].runs:
            set_run_font(run, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5)
    set_table_widths(table, widths)
    doc.add_paragraph()


def english_doc():
    doc = Document()
    configure_document(doc, "UIWEB Project Technical Guide")
    add_title_page(
        doc,
        "UIWEB Project Technical Guide",
        "Architecture, workflows, feature history, code boards, and future direction",
        "English version",
        "AdminSite API, AdminSite Frontend, UserSite, SharedComponents, Contracts, Resource Manager, Content Management, Page Builder, tools, and operational workflows.",
    )

    add_h1(doc, "1. Executive Snapshot")
    add_para(doc, "UIWEB is a three-application CMS and website platform. The AdminSite is where authenticated admin users build pages, manage content, manage resources, configure branding and theme, and review operational data. The UserSite is the public-facing website. The API is the shared backend that owns authentication, MongoDB persistence, publish/reset workflows, asset storage, content workflow, resource management, public rendering data, forms, and metrics.")
    add_para(doc, "The core architectural idea is separation between draft editing and public output. Admin users edit draft collections. Publishing clones a consistent page graph or content item into published collections. UserSite reads only published data through public endpoints, so the public site does not accidentally expose draft work.")
    add_para(doc, "The second important idea is separation between lightweight direct uploads and reusable managed resources. Direct uploads stay attached to a page, section, block, branding field, or content field. Managed resources live in Resource Library and can be selected across pages and content. Deleting or replacing assets is routed through cleanup and usage checks so unused storage objects can be removed while used assets remain protected.")
    add_bullets(doc, [
        "AdminSite-API is the system authority for business rules and persistence.",
        "AdminSite-Frontend is a Blazor Server dashboard for editors, managers, and AdminAdmin users.",
        "UserSite is a Blazor Server public site that consumes public API DTOs.",
        "SharedComponents contains the renderer used by both public pages and admin previews.",
        "Contracts holds DTOs shared across API, AdminSite, UserSite, and SharedComponents.",
        "MongoDB stores structured data and asset metadata. Uploaded binary files currently live in R2-compatible object storage, not directly inside MongoDB.",
    ])

    add_code_board(doc, "System Board", """
[Admin User]
    |
    v
AdminSite-Frontend (Blazor Server dashboard)
    |
    v
AdminSite-API (ASP.NET Core Web API)
    |
    +--> MongoDB collections: draft, published, users, resources, forms, settings
    |
    +--> R2StorageService: uploaded images, files, videos
    |
    v
Public API endpoints under /api/public
    |
    v
UserSite + SharedComponents renderer
    |
    v
Public website visitors
""")

    add_h1(doc, "2. Repository and Module Map")
    add_table(doc, ["Folder", "Responsibility", "Important Notes"], [
        ("AdminSite-API", "ASP.NET Core backend. Owns controllers, services, MongoDB collections, authentication, publish/reset, content workflow, resource management, public API, forms, metrics, and asset storage.", "Project file: AdminSite-API/Main-API.csproj. Root namespace is FullProject."),
        ("AdminSite-Frontend", "Blazor Server admin dashboard. Owns UI screens for Page Builder, Content Management, Resource Library, global settings, users, forms, and metrics.", "Uses HttpService to call API and AdminAuthService to keep the admin session in localStorage."),
        ("UserSite", "Public Blazor Server website. Loads public navigation, theme, pages, content pages, footer, social buttons, and global buttons from API.", "Uses PublicApiService and SharedComponents/PageRenderer.razor."),
        ("SharedComponents", "Shared public renderer and visual section/block components.", "Important because Admin preview and UserSite should render the same data shape."),
        ("Contracts", "Shared DTO contracts and global helper DTOs.", "Contracts/Public drives renderer DTOs; Contracts/Auth defines roles and permissions."),
        ("Tool", "Actual project tools used by the owner/testers.", "Contains demo database import tool and CloneUtility coverage tool."),
        ("AI-Tools", "Helper scripts and reports mainly used by Codex during audits.", "Separate from user-facing project tools for clarity."),
        ("build-check", "Local build output folder used to avoid locked Visual Studio DLLs.", "Should remain local/ignored."),
    ], [1900, 4560, 2900])

    add_code_board(doc, "Runtime Application Board", """
AdminSite-Frontend
    Pages/*.razor, Services/*.cs
    -> calls API via HttpService
    -> stores admin_session in browser localStorage

AdminSite-API
    Controllers/*.cs
    -> Services/*.cs
    -> MongoDbContext collections
    -> R2StorageService for binary assets

UserSite
    Components/Pages/PageView.razor
    -> PublicApiService
    -> SharedComponents/PageRenderer.razor
    -> SharedComponents/Sections and Blocks
""")

    add_h1(doc, "3. Startup and Configuration Workflow")
    add_para(doc, "The API startup lives in AdminSite-API/Program.cs. It reads MongoDb, Jwt, Seed, Cors, R2Storage, and FormSecurity sections from appsettings.json. If MongoDB settings are missing, the app fails fast. If JWT secret is missing or too short, the app also fails fast. This protects the backend from running with an invalid authentication configuration.")
    add_para(doc, "Services are registered by feature area: auth, pages, sections, blocks, publish/reset, content services, managed resource services, form services, public services, metrics, settings, asset cleanup, and R2 storage. MongoClient and IMongoDatabase are singletons, while most application services are scoped.")
    add_para(doc, "On startup the API attempts to create MongoDB indexes by calling MongoIndexService.EnsureIndexesAsync. Index creation failures are logged as warnings, and the app continues. This is intentional because a temporary index issue should not fully block local testing, but slow queries may happen until indexes are fixed.")
    add_para(doc, "Startup also removes legacy default Resource Library albums and ensures default public form definitions exist. Admin seeding is optional and only happens if Seed:AdminEmail and Seed:AdminPassword are configured.")
    add_code_board(doc, "API Startup Board", """
Program.cs
    |
    +--> read appsettings: MongoDb, Jwt, Cors, R2Storage, FormSecurity
    |
    +--> register MongoClient + IMongoDatabase + MongoDbContext
    |
    +--> register services by feature module
    |
    +--> configure JWT authentication + fallback authorization
    |
    +--> configure rate limits: public-form, admin-login
    |
    +--> EnsureIndexesAsync()
    |
    +--> RemoveLegacyDefaultAlbumsAsync()
    |
    +--> EnsureDefaultDefinitionsAsync()
    |
    +--> optional SeedAdminAsync()
    |
    v
Middleware: CORS -> HTTPS -> RateLimiter -> Authentication -> AdminSessionValidationMiddleware -> Authorization -> Controllers
""")

    add_h1(doc, "4. MongoDB Data Model")
    add_para(doc, "MongoDbContext exposes the main collections. The most important structural decision is draft/published separation. Pages, sections, blocks, and content items are edited in draft collections and copied into published collections when published. The public website reads published collections only.")
    add_table(doc, ["Collection Group", "Collections", "Purpose"], [
        ("Page Builder", "pages_draft, pages_published, page_revisions", "Stores page metadata, slug tree, SEO, page card, draft state, and published snapshots."),
        ("Sections and Blocks", "sections_draft, sections_published, blocks_draft, blocks_published", "Stores page layout sections and reusable block content."),
        ("Content", "content_draft, content_published, content_types, content_revisions, content_audit_logs", "Stores articles, whitepapers, video resources, image resources, gallery resources, and their workflow history."),
        ("Resource Library", "managed_resources, resource_albums", "Stores reusable asset metadata, album organization, source type, storage key, and usage-protected resource records."),
        ("Admin Auth", "admin_users, admin_sessions, admin_login_activity, admin_audit_logs", "Stores admin accounts, active/revoked sessions, login records, and audit records."),
        ("Forms", "form_definitions, form_submissions", "Stores reusable form definitions and public submissions."),
        ("Global Site", "branding, site_settings, theme, footer, social, global_buttons", "Stores global visual and navigation configuration. Some are accessed through services by collection name."),
        ("Metrics", "visitor_metrics", "Stores page/content/download counters for Website Activity."),
    ], [1750, 3400, 4210])
    add_para(doc, "MongoDB stores metadata and JSON-like content structures. It does not store uploaded image/video/file bytes in the normal application path. Current uploaded files go to R2 through R2StorageService; MongoDB stores URL, storage key, file name, MIME type, size, source, and resource identity.")

    add_h1(doc, "5. Authentication and Authorization")
    add_para(doc, "Authentication is built around JWT plus server-side session validation. A valid JWT is not enough by itself: the token must also match an active admin session stored in MongoDB. This gives the system a way to revoke sessions after logout, password reset, role changes, permission changes, disabling a user, or deleting an account.")
    add_para(doc, "Roles are defined in Contracts/Auth/AuthDtos.cs. The current roles are AdminAdmin, Manager, Writer, and Viewer. Permissions are string constants: page-builder, manage-content, publish-content, manage-users, manage-settings, delete-content, and view-logs. AdminAdmin is the full-access role and bypasses permission checks through AdminAuthorization.HasPermission.")
    add_para(doc, "AuthService.LoginAsync normalizes email, checks account status, applies lockout rules, verifies password with BCrypt, creates a token id, builds JWT claims, stores an AdminSessionRecord, resets failed login counters, and writes login/audit records. Failed logins increment FailedLoginAttempts and can lock the account for fifteen minutes after five failed attempts.")
    add_para(doc, "AdminSessionValidationMiddleware runs after JWT authentication. It skips public endpoints, Swagger, login, and health. For other requests it reads adminId, jti, and tokenVersion from JWT claims, then calls AuthService.ValidateSessionAsync. Invalid sessions return 401 and X-Admin-Session-Invalid: true. The AdminSite HttpService treats that header as a real session expiry and navigates the user back to /login.")
    add_code_board(doc, "Authentication Board", """
Login.razor
    |
    v
AuthController.Login
    |
    v
AuthService.LoginAsync
    |
    +--> admin_users: find email, status, lockout, password hash
    +--> admin_sessions: insert session with token id + token version
    +--> admin_login_activity/admin_audit_logs: record event
    |
    v
JWT returned to AdminSite
    |
    v
AdminAuthService stores admin_session in localStorage
    |
    v
HttpService attaches Bearer token to API requests
    |
    v
JWT middleware validates token signature/lifetime
    |
    v
AdminSessionValidationMiddleware validates MongoDB session
    |
    v
Controllers apply role/permission checks
""")
    add_table(doc, ["Role", "Default Capability", "Important Limit"], [
        ("AdminAdmin", "Full access to users, settings, page builder, content, publish/delete, logs, and Resource Library.", "Cannot remove the last active AdminAdmin account."),
        ("Manager", "Content management, publish content, delete content by default.", "User management and settings require explicit permission."),
        ("Writer", "Manage content and Resource Library access. Can create/edit content according to UI rules.", "Does not automatically publish or manage users."),
        ("Viewer", "Read-only dashboard role.", "Should not view Resource Library according to the current Resource Manager direction."),
    ], [1700, 4800, 2860])

    add_h1(doc, "6. Page Builder Workflow")
    add_para(doc, "The Page Builder is built around Page, Section, and Block documents. Pages define the navigation and identity of a public page. Sections define the vertical page structure. Blocks define reusable content units inside sections or columns. Draft records are edited in AdminSite; published records are created only by PublishService.")
    add_para(doc, "The most important identity fields are StableId, SourceId, Version, and PublishedAt. StableId is the long-term logical identity across draft and published versions. SourceId records which draft document produced a published snapshot. Version increments when publishing. Public rendering usually tracks PageStableId and SectionStableId, not just MongoDB document ids.")
    add_para(doc, "ColumnsSection requires special care. Column slots have IDs, and blocks inside columns point to those slot IDs through ColumnSlotId. CloneUtility must preserve ColumnSlot.Id and Block.ColumnSlotId during publish/reset, otherwise blocks can render in the wrong column or disappear from columns.")
    add_code_board(doc, "Page Builder Edit and Publish Board", """
Canvas.razor / EditPanel.razor
    |
    +--> AdminPageService -> PagesController -> PageService -> pages_draft
    +--> AdminSectionService -> SectionsController -> SectionService -> sections_draft
    +--> AdminBlockService -> BlocksController -> BlockService -> blocks_draft
    |
    v
User clicks Publish
    |
    v
PagesController.POST /api/admin/pages/{pageId}/publish
    |
    v
PublishService.PublishPageAsync
    |
    +--> load draft page, sections, blocks
    +--> delete existing published page graph for same StableId
    +--> CloneUtility.ClonePage / CloneSection / CloneBlock
    +--> insert into pages_published, sections_published, blocks_published
    +--> mark draft page Published and increment Version
    +--> after commit: AssetCleanupService removes unused replaced assets
""")
    add_table(doc, ["File", "Responsibility", "Why It Matters"], [
        ("AdminSite-API/Controllers/PagesController.cs", "Admin page endpoints for CRUD, publish, reset, and page hierarchy operations.", "Entry point for most Page Builder commands."),
        ("AdminSite-API/Services/PagesServices.cs", "Page persistence, slug/full slug logic, page retrieval, children, and public page lookup.", "Keeps page identity and navigation consistent."),
        ("AdminSite-API/Controllers/SectionsController.cs", "Section endpoints.", "Connects UI section editors to SectionService."),
        ("AdminSite-API/Services/SectionServices/SectionServices.cs", "Section creation, update, ordering, visibility, and type-specific behavior.", "Owns the data shape of page sections."),
        ("AdminSite-API/Services/BlockService.cs", "Block creation/update/delete/order and public block retrieval.", "Owns block data and column/container relationships."),
        ("AdminSite-API/Services/PublishAndResetService/PublishService.cs", "Draft-to-published transaction.", "Creates the public snapshot."),
        ("AdminSite-API/Services/PublishAndResetService/ResetService.cs", "Published-to-draft reset transaction.", "Restores draft from public snapshot."),
        ("AdminSite-API/Utils/CloneUtility.cs", "Manual deep clone of page, section, block graph.", "Current design debt because new fields can be missed silently."),
    ], [3100, 3300, 2960])

    add_h1(doc, "7. CloneUtility and the Current Design Weakness")
    add_para(doc, "CloneUtility currently performs manual per-type copying. It knows every section type and every block type, then constructs a new object by copying fields one by one. This works when maintained carefully, but it is fragile: if a future developer adds a new field to a model and forgets to add it to CloneUtility, publishing or reset can silently drop that field.")
    add_para(doc, "One real weakness already found was around ColumnSlotId. ColumnsSection uses slot IDs and blocks reference those slots. If the clone operation does not preserve those IDs, a published page can look different from the draft page. Recent cleanup improved coverage and protected the column slot behavior, but the broader design should still be replaced.")
    add_para(doc, "The intended future design is a PageGraphCloneService or clone-profile system. Instead of one manual switch statement being responsible for every model field forever, the clone should use a serializer/profile approach with explicit patch steps for identity fields. Profiles could be PublishSnapshot, DraftResetSnapshot, PresetApply, and DuplicateAsNewContent. That would make new fields safer because default deep copy would keep them unless a profile explicitly changes them.")
    add_code_board(doc, "Future Clone Design Board", """
Current:
    CloneUtility.CloneSection(section)
        switch on concrete type
        manually copy every field
        risk: new field can be forgotten

Better:
    PageGraphCloneService.Clone(graph, profile)
        deep-copy by serialization or mapping profile
        apply identity patch rules:
            - new MongoDB Id
            - preserve or create StableId based on profile
            - set SourceId
            - set PublishedAt
            - set Status
            - set Version
        run coverage/audit for fields excluded by profile
""")

    add_h1(doc, "8. Content Management Workflow")
    add_para(doc, "Content Management is separate from Page Builder. It manages article-like pages and resource-like content records such as whitepapers, videos, images, and galleries. The major design improvement was adding content behavior. A content type can be Page, FileResource, VideoResource, ImageResource, or Gallery. Behavior controls editor fields, validation, preview behavior, and LibrarySection rendering.")
    add_para(doc, "ContentService is now a facade over smaller services. The user decided not to do a partial/half split of the entire service too early, but the important behavior modules already exist. ContentTypeService owns type CRUD and defaults. ContentValidationService owns behavior-aware validation. ContentWorkflowService owns submit/publish/delete/restore/permanent delete. ContentRevisionService owns revisions and audit logs. ContentMappingService maps models to responses and published clones. ContentAssetMetadataService normalizes ResourceId, ResourceSource, StorageKey, body media, gallery items, and attachments.")
    add_para(doc, "Page behavior requires page-like body content. FileResource requires a file. VideoResource requires a video. ImageResource requires an image. Gallery requires gallery items. The current strict rule is that Content Page Preview should render Page behavior only. Resource-based content is monitored through All Content and My Content lists and through Resource/Library flows, not by pretending every resource is an article page.")
    add_code_board(doc, "Content Create/Update Board", """
ContentEditorShell.razor / ContentItems.razor
    |
    v
AdminContentService
    |
    v
ContentController
    |
    v
ContentService
    |
    +--> ContentValidationService: behavior-aware rules
    +--> ContentAssetMetadataService: ResourceId/ResourceSource/StorageKey normalization
    +--> ContentRevisionService: save revision before update
    +--> MongoDB content_draft
    +--> AssetCleanupService: delete replaced direct-upload assets if unused
""")
    add_code_board(doc, "Content Publish/Delete Board", """
ContentController.Publish/Delete/Restore/PermanentDelete
    |
    v
ContentWorkflowService
    |
    +--> Start MongoDB transaction
    +--> Validate publish rules
    +--> Update content_draft status
    +--> CloneForPublished via ContentMappingService
    +--> Replace content_published record by StableId
    +--> Write revision/audit log
    +--> Commit transaction
    |
    v
AssetCleanupService deletes unused replaced assets after commit
""")
    add_table(doc, ["Behavior", "Editor Rule", "Preview/Display Rule"], [
        ("Page", "Title, summary, cover/hero, body blocks/html are enabled.", "Content Page Preview and public content detail page."),
        ("FileResource", "Managed file selector/upload is required; body is hidden.", "LibrarySection opens file in new tab; All/My Content list link opens file directly."),
        ("VideoResource", "Video upload or valid YouTube video link is required; body is hidden.", "LibrarySection opens video modal/player."),
        ("ImageResource", "Image selector/upload is required; body is hidden.", "LibrarySection opens image lightbox/modal."),
        ("Gallery", "Gallery items required eventually; body is hidden.", "LibrarySection Gallery layout shows horizontal/gallery display."),
    ], [1750, 4010, 3600])

    add_h1(doc, "9. Resource Manager and Resource Library")
    add_para(doc, "Resource Manager was overhauled to become a true library, not a second content page preview. It manages reusable assets: images, files, and videos. Content Management chooses selected files or media for public display, while Resource Library owns the larger reusable pool. Those functions are related but not the same.")
    add_para(doc, "The current rule is one resource, one album at most. Albums are organizational metadata only. They do not restrict where a resource can be reused. A resource can be used by multiple pages, sections, blocks, branding fields, and content records even if it belongs to a single album.")
    add_para(doc, "Albums are scope separated. Media albums can contain images and videos. File albums can contain files. Mixed page-composition groups are intentionally handled by naming convention rather than mixed albums, for example File-Page1 and Resource-Page1. The delete rule for albums is hard-block: if an album has resources, the user must move or unfile those resources before deleting the album.")
    add_para(doc, "The Resource Library UI direction is a central list/grid gallery with a compact sidebar. The main area is for browsing and selecting resources. The sidebar is for details: preview or file icon, name, description, Preview action, usage, file size, MIME type, uploaded by, created time, and updated time. Internal fields like storage key and raw URL should not be shown to non-code users.")
    add_code_board(doc, "Resource Library Board", """
ContentResources.razor
    |
    +--> Normal mode: resource list/grid, type tabs, search, view switch
    +--> Album mode: folder-like albums, scoped to media or file
    +--> Add Resource: type-first modal, upload queue, optional YouTube link for video
    +--> Add Album/Group: create organizational folder
    |
    v
AdminContentService
    |
    v
ManagedResourcesController
    |
    +--> ManagedResourceService
        +--> ManagedResourceValidationService
        +--> ManagedResourceAlbumService
        +--> ManagedResourceUsageService
        +--> AssetCleanupService
""")
    add_code_board(doc, "Resource Delete Board", """
User clicks Delete Resource
    |
    v
ManagedResourcesController.Delete
    |
    v
ManagedResourceService.DeleteAsync
    |
    +--> ManagedResourceUsageService.GetUsageAsync(resource)
    |
    +--> if UsageCount > 0:
    |       return error: remove references before deleting
    |
    +--> if UsageCount == 0:
            delete managed_resources record
            AssetCleanupService.DeleteUnusedAsync(resource.Url, resource.ThumbnailUrl)
            AssetReferenceService verifies no remaining references
            R2StorageService deletes storage object
""")
    add_table(doc, ["Resource Type", "Storage Rule", "UI Rule"], [
        ("Image", "Uploaded binary goes to R2/local storage abstraction. Mongo stores URL, StorageKey, file metadata, album id, and usage metadata.", "Shown as image list/grid. Choose from Resource Library or upload new."),
        ("File", "Uploaded binary goes to storage. Mongo stores URL, StorageKey, filename, MIME, size.", "Grid uses file icon/placeholder by extension. No editable preview image."),
        ("Video Upload", "Uploaded video file goes to storage. Mongo stores URL, StorageKey, filename, MIME, size. Accepted formats include mp4, webm, mov.", "No custom thumbnail editing. Use video label/name for navigation."),
        ("Video YouTube Link", "Video-only exception. No R2 upload. Mongo stores normalized YouTube embed URL and Source = external-url.", "Allowed only for real YouTube video links, not playlist/channel/search pages."),
    ], [1800, 4700, 2860])

    add_h1(doc, "10. Resource Picker and Asset Metadata")
    add_para(doc, "Resource Picker should save metadata, not just copy a URL. The important fields are ResourceId, ResourceSource, and StorageKey, along with the visible URL. This is what lets the system later understand whether an asset is a DirectUpload or a ManagedResource, track usage, propagate replacements, and clean unused old assets safely.")
    add_para(doc, "DirectUpload means the asset is attached to a specific field and does not appear in Resource Library. ManagedResource means the asset is a reusable library item. Hero backgrounds and decorative section/block images usually default to DirectUpload. Whitepapers, reports, tool files, gallery items, and reusable video/image resources default to ManagedResource.")
    add_code_board(doc, "Picker Metadata Board", """
ResourcePickerModal.razor
    |
    v
User selects resource
    |
    v
Target editor field receives:
    Url
    ResourceId
    ResourceSource = ManagedResource
    StorageKey
    FileName / Mime / Size when relevant
    |
    v
ContentAssetMetadataService or section/block update mapper normalizes values
    |
    v
MongoDB stores both visible URL and source metadata
""")

    add_h1(doc, "11. Asset Storage and Cleanup")
    add_para(doc, "The project started with base64-style asset storage ideas, but that made documents and database records too large. The current design stores binary assets outside MongoDB through R2StorageService. MongoDB stores references and metadata. This is also the right design for a future company-local storage backend: replace the storage implementation while keeping the same metadata contract.")
    add_para(doc, "AssetCleanupService is the central cleanup rule. It receives old URLs from replaced or deleted records, asks AssetReferenceService whether those URLs are still referenced anywhere, and only then calls R2StorageService.DeleteAsync. This protects shared resources and avoids deleting files that are still used.")
    add_para(doc, "Resource Library deletion follows the same global rule. If a managed resource is used anywhere, it cannot be deleted. If it is unused and the user chooses to erase it, the record is deleted and the storage object is cleaned. Direct-upload assets are cleaned when replaced or when their owning page/content graph is deleted, as long as no other record references the same URL.")
    add_code_board(doc, "Asset Cleanup Board", """
Old asset URL detected
    |
    v
AssetCleanupService.DeleteIfUnusedAsync(oldUrl, replacementUrl)
    |
    +--> skip if oldUrl is empty
    +--> skip if oldUrl == replacementUrl
    +--> AssetReferenceService.IsReferencedAsync(oldUrl)
    |
    +--> if referenced:
    |       keep storage object
    |
    +--> if not referenced:
            R2StorageService.DeleteAsync(oldUrl)
""")
    add_table(doc, ["Storage Layer", "Current Implementation", "Future Company-Local Direction"], [
        ("Binary storage", "Cloudflare R2/S3-compatible storage through R2StorageService.", "Implement equivalent local storage service that returns URL/path and storage key."),
        ("Database", "MongoDB stores metadata, not bytes: URL, StorageKey, ResourceId, ResourceSource, MIME, size.", "Keep the same metadata model if possible to avoid rewriting editors/renderers."),
        ("Cleanup", "AssetCleanupService and AssetReferenceService protect used assets.", "Reuse the same cleanup service against company-local storage."),
        ("Security", "Extension, MIME, folder, and signature checks exist. Anti-virus/deep scan is future work.", "Add scan/quarantine pipeline before final storage persistence."),
    ], [1800, 3700, 3860])

    add_h1(doc, "12. Public Rendering and UserSite")
    add_para(doc, "The public website does not read admin draft collections directly. UserSite calls /api/public endpoints. PublicController delegates page assembly to PublicPageAssemblyService and metadata to PublicMetadataService. The response is mapped to shared public DTOs, then rendered by SharedComponents/PageRenderer.razor.")
    add_para(doc, "PublicPageAssemblyService fetches a published page, visible published sections, and visible published blocks. It handles special sections first. ShowcaseSection reads published child pages and maps them into cards. LibrarySection reads published content items according to selected content types and maps them into resource-aware library items. ColumnsSection groups blocks by ColumnSlotId. Standard sections are mapped through a catch-all branch.")
    add_para(doc, "UserSite/PublicApiService receives JSON, maps section types into Contracts.Public DTO classes, builds block trees, maps column slot blocks, and passes the final PublicPageDto to PageRenderer. PageRenderer pattern matches each PublicSectionDto subtype and calls the matching component under SharedComponents/Sections.")
    add_code_board(doc, "Public Page Render Board", """
Visitor navigates to /solutions
    |
    v
UserSite PageView.razor
    |
    v
PublicApiService.GetPageAsync("solutions")
    |
    v
GET AdminSite-API /api/public/pages/solutions
    |
    v
PublicController.GetPage
    |
    v
PublicPageAssemblyService.GetPageResponseAsync
    |
    +--> PageService.GetByFullSlugAsync
    +--> SectionService.GetPublicSectionsByPageAsync
    +--> BlockService.GetPublicByPageAsync
    +--> map special sections and standard sections
    |
    v
Contracts.Public DTOs
    |
    v
SharedComponents/PageRenderer.razor
    |
    v
HeroSection, LibrarySection, ColumnsSection, etc.
""")

    add_h1(doc, "13. LibrarySection Rendering")
    add_para(doc, "LibrarySection became the main renderer for managed content resources. It supports existing layouts: Card, Grid, Rows, List, and Gallery. Resource behavior controls what the call-to-action does. FileResource opens the file in a new browser tab. VideoResource opens a video modal/player. ImageResource opens a lightbox/modal. Gallery shows a gallery-style layout and opens images/videos in modal/lightbox views.")
    add_para(doc, "This replaced the old overlap with GallerySection. The direction was to add a Library Gallery layout first, keep old renderers only for old pages if needed, hide GallerySection from Add Section, and then clean it up after migration. Resource-based content should not pretend to be normal article pages.")
    add_code_board(doc, "LibrarySection Behavior Board", """
LibrarySectionEditor.razor
    |
    v
LibrarySection stored in sections_draft/published
    |
    v
PublicPageAssemblyService detects LibrarySection
    |
    v
ContentService.GetPublishedLibraryItemsAsync
    |
    v
MapLibraryItem(content, contentType)
    |
    v
SharedComponents/Sections/LibrarySection.razor
    |
    +--> Page content: link to content detail page
    +--> FileResource: open file URL in new tab
    +--> VideoResource: modal video player
    +--> ImageResource: lightbox/modal
    +--> Gallery: gallery display with modal/lightbox
""")

    add_h1(doc, "14. Forms and Public Submission Security")
    add_para(doc, "Forms are stored as FormDefinition records and can be used by FormBlock or modal forms. Startup calls FormDefinitionService.EnsureDefaultDefinitionsAsync, so common public modal forms are available without manual setup. Public submissions go through public endpoints and are protected by rate limiting and request size limits.")
    add_para(doc, "PublicController exposes form submit endpoints under page/section/block paths and modal form paths. Those endpoints are AllowAnonymous because public visitors submit them, but they use the public-form rate limiting policy. FormSubmissionSecurityService and FormValidationService handle validation and abuse checks such as honeypot and payload limits.")
    add_code_board(doc, "Public Form Board", """
Visitor submits form
    |
    v
SharedComponents/Blocks/FormBlock.razor
    |
    v
UserSite PublicApiService.SubmitFormAsync
    |
    v
POST /api/public/pages/{slug}/sections/{sectionId}/blocks/{blockId}/form/submit
    |
    v
PublicFormSubmissionHandler
    |
    +--> FormDefinitionService
    +--> FormValidationService
    +--> FormSubmissionSecurityService
    +--> FormSubmissionService
    |
    v
MongoDB form_submissions
""")

    add_h1(doc, "15. Theme, Branding, Navigation, and Global Modules")
    add_para(doc, "Global modules configure the parts of the public site that are not tied to one page body: branding, theme, navigation, footer, social buttons, and global buttons. PublicMetadataService exposes these through /api/public endpoints. UserSite loads them through PublicApiService and uses them in layout components such as NavMenu, MainLayout, SiteHeader, SiteFooter, and SocialFloat.")
    add_para(doc, "Recent direction for Theme and UserSite layout is that the top page navigation bar should derive colors from Theme rather than hard-coded static colors. Sections also gained or should keep a Background Type named Theme so a section can automatically use theme color. Alignment fixes were made/expected for the nav items so they sit visually centered in the bar.")
    add_code_board(doc, "Global Metadata Board", """
Admin Global UI
    |
    +--> Branding.razor -> BrandingController -> BrandingService
    +--> Theme.razor -> ThemeController -> ThemeService
    +--> Footer.razor -> FooterController -> FooterService
    +--> Social.razor -> SocialController -> SocialButtonsService
    +--> GlobalButtons.razor -> GlobalButtonsController -> GlobalButtonsService
    |
    v
MongoDB global collections/settings
    |
    v
PublicMetadataService
    |
    v
/api/public/branding, /theme, /footer, /social, /global-buttons, /navigation
    |
    v
UserSite layout + SharedComponents
""")

    add_h1(doc, "16. Upload Security")
    add_para(doc, "UploadSecurityPolicy currently checks allowed upload folders, file extensions, MIME types, and file signatures. For managed resources it checks that the selected kind matches the configured allowed formats. Video signatures include mp4/mov ftyp and webm signatures. Office files, PDF, images, and text files are also signature checked.")
    add_para(doc, "This is useful baseline protection, but it is not anti-virus scanning. The future plan is to add anti-virus/deep scanning and database-bomb prevention. The safest future design is an upload quarantine pipeline: receive file, validate extension/MIME/signature/size, scan file, store only after scan passes, then create or update the database record.")
    add_code_board(doc, "Future Upload Scan Board", """
Upload request
    |
    v
Controller request size + selected kind
    |
    v
UploadSecurityPolicy: folder + extension + MIME + signature
    |
    v
Quarantine temporary storage
    |
    v
Anti-virus / deep scan / archive-bomb checks
    |
    +--> fail: delete quarantine file, return error
    |
    +--> pass: move to permanent storage, save Mongo metadata
""")

    add_h1(doc, "17. Metrics and Website Activity")
    add_para(doc, "VisitorMetricService increments counters for public page views, public content page views, and downloads. PublicController increments metrics after successful page/content lookup, and exposes a small download tracking endpoint. AdminSite page WebsiteActivity.razor displays these metrics for admin review.")
    add_para(doc, "The metrics system is intentionally simple. It is useful for seeing what pages and resources are active without building a full analytics platform. If the project later requires advanced analytics, the current counter model can remain as a lightweight internal signal while a dedicated analytics tool handles deeper visitor behavior.")

    add_h1(doc, "18. Demo Database Import Tool")
    add_para(doc, "The project includes a one-click import tool under Tool/Tool For Demo Import. It is intentionally import-only. It imports prepared JSON collections into a MongoDB database so a new local tester can reproduce the current UserSite layout and admin demo data without needing the owner's personal database.")
    add_para(doc, "The tool can create the target database when it imports. The discussed target name is FullProjectDb-UIWEB-3. Testers do not need to manually create that database first. They do need to configure the application appsettings.json to point to that database when they want to test against the imported demo.")
    add_para(doc, "The import tool does not duplicate R2 files. It imports metadata and URLs. If the future company-local storage migration happens, a separate migration/import project should be created after the company storage details are known.")
    add_code_board(doc, "Demo Import Board", """
run-import.bat
    |
    v
DemoDbImporter.exe / DemoDbImporter project
    |
    +--> appsettings.importer.json: MongoDB connection + target database
    +--> demo-seed/manifest.json: collection order
    +--> demo-seed/collections/*.json: data snapshots
    |
    v
MongoDB target database, for example FullProjectDb-UIWEB-3
    |
    v
AdminSite-API appsettings.json must point to that database for testing
""")

    add_h1(doc, "19. Project History and Direction")
    add_para(doc, "The project evolved through several major decisions. Early asset ideas used base64-like data, but that would make the database too large. The system moved toward external storage with URL/storage key metadata. Resource Management was then introduced so not every upload becomes a reusable resource. This preserved two clean paths: DirectUpload for one-off assets and ManagedResource for reusable library assets.")
    add_para(doc, "The LibraryFeature plan originally had nine phases: define asset paths, content type behavior, Resource Management base, Resource Picker, Content Editor integration, LibrarySection integration, Media/Resource Preview, GallerySection deprecation, and safety cleanup. A later Phase 10 was service split/cleanup, including ContentService decomposition. The user explicitly preferred real service splits over partial cleanup.")
    add_para(doc, "Resource Manager Overhaul 2 later refined Resource Library into a professional library UI. The main list/grid became the center of the experience. The sidebar became details and usage. Albums became folder-like organization, one resource per album, no automatic Unsorted album, hard-block delete if album contains resources, and type-first upload modals with a better upload queue.")
    add_para(doc, "The current direction is stability and clarity. Major themes are: keep phase order, avoid overlapping screens, hide code/internal fields from non-code admins, protect assets from accidental deletion, clean unused storage objects, keep Resource Library focused on reusable assets, and eventually redesign CloneUtility so publishing cannot silently miss new fields.")
    add_table(doc, ["Period", "What Changed", "Why It Matters"], [
        ("Early asset direction", "Base64-style storage was rejected in favor of URL/storage key metadata.", "Prevents MongoDB bloat and keeps storage replaceable."),
        ("LibraryFeature", "DirectUpload vs ManagedResource, behavior-driven content, picker metadata, LibrarySection rendering.", "Created clear separation between one-off uploads and reusable resources."),
        ("Resource Manager Overhaul 2", "Central gallery/list UI, compact sidebar, albums, bulk upload queue, resource settings, role visibility.", "Made Resource Library understandable for non-code users."),
        ("Content service cleanup", "ContentService facade plus type, validation, workflow, revision, mapping, asset metadata services.", "Reduced one large service without doing a fake partial split."),
        ("Asset cleanup", "Central AssetCleanupService verifies references before storage deletion.", "Stops storage bloat while protecting reused assets."),
        ("Demo import tool", "Import-only tool for local demo database setup.", "Makes local/company testing easier without exporting from personal DB each time."),
        ("Current stability work", "Mongo indexes, CloneUtility coverage, workflow transactions, namespace/csproj consistency.", "Hardens the codebase before larger migrations."),
    ], [1850, 4010, 3500])

    add_h1(doc, "20. Current Known Weaknesses and Future Work")
    add_bullets(doc, [
        "CloneUtility should be redesigned into a safer PageGraphCloneService or clone profile system.",
        "Log Overhaul is still planned but intentionally not touched during recent stability phases. The direction is separate Login Log and Audit Trail, with retention/archive/export instead of casual deletion.",
        "Company-local storage migration needs more information about the company database/storage environment. R2StorageService should be replaced behind the same abstraction rather than rewriting business modules.",
        "Anti-virus/deep scanning/database-bomb prevention should be added to upload flow.",
        "Models.cs and admin DTO/model splitting should be done later as a real structural cleanup, not as a rushed partial split.",
        "Resource Library settings should remain admin-governed and should not expose confusing internal storage fields to normal users.",
        "Testing should grow around publish/reset, CloneUtility field coverage, Resource Library deletion/usage, content behavior validation, and asset cleanup.",
        "Tool folder ownership should remain clear: Tool for actual project tools, AI-Tools for assistant/audit helper scripts.",
    ])

    add_h1(doc, "21. Developer Orientation: How to Change the System Safely")
    add_para(doc, "When adding a new field to Page, Section, Block, ContentItem, ManagedResource, or DTOs, do not only update the UI. Check the whole lifecycle: create/update DTOs, server model, validation, mapping, clone/publish/reset, public DTO mapping, renderer, resource replacement propagation, asset cleanup, database importer seed, and tests.")
    add_code_board(doc, "New Field Safety Checklist Board", """
New field added
    |
    +--> Model: AdminSite-API/Models.cs
    +--> DTO: AdminSite-API/DTOs.cs or Contracts/*
    +--> Admin UI model: AdminSite-Frontend/Models/AdminModels.cs
    +--> Editor UI and service mapper
    +--> Controller create/update endpoint
    +--> Validation service
    +--> Publish/reset clone or content mapping
    +--> PublicPageAssemblyService mapping if public
    +--> Contracts.Public DTO if rendered
    +--> SharedComponents renderer if visible
    +--> AssetCleanupService if field stores asset URL
    +--> ManagedResourceService replacement propagation if field can use Resource Library
    +--> Demo importer seed if demo data requires it
    +--> Focused test or coverage tool
""")

    add_h1(doc, "22. Appendix: High-Value Code Boards")
    add_code_board(doc, "All Content File Link Behavior", """
All Content / My Content list row click
    |
    +--> If content behavior == Page:
    |       open content page/detail route
    |
    +--> If content behavior == FileResource:
    |       open file URL directly in new tab
    |
    +--> If content behavior == VideoResource/ImageResource/Gallery:
            use resource-aware action rather than article-page redirect
""")
    add_code_board(doc, "Video Background Direction", """
SectionStylePanel.razor
    |
    +--> Background Type: Video
    +--> Upload New Video
    +--> Choose From Resource Library
    |
    v
Section.Style.BackgroundVideoUrl + metadata
    |
    v
SharedComponents/Helpers/StyleHelper.cs
    |
    v
SectionBackgroundMedia.razor renders muted autoplay loop video

YouTube/external URLs remain for Video Blocks or content video fields, not section background.
""")
    add_code_board(doc, "Background Image Fit Direction", """
Section background image
    |
    +--> Fit = Cover
    |       fills section, may crop image
    |
    +--> Fit = Contain
    |       shows full image, may leave empty space
    |
    +--> Position = center/top/bottom/left/right
            controls visible focal area
""")
    add_code_board(doc, "Resource Replacement Propagation", """
Resource Editor: Replace File
    |
    v
ManagedResourceService.ReplaceUploadAsync
    |
    +--> validate same kind
    +--> update ManagedResource URL/storage/file metadata
    +--> propagate to:
            content_draft/content_published
            pages_draft/pages_published
            sections_draft/sections_published
            blocks_draft/blocks_published
            branding
    +--> AssetCleanupService.DeleteIfUnusedAsync(oldUrl, newUrl)
""")
    add_code_board(doc, "Admin Request Lifecycle", """
AdminSite UI event
    |
    v
Admin service class
    |
    v
HttpService.SendAsync
    |
    +--> attach JWT from localStorage
    +--> call API
    +--> parse ApiResponse<T>
    +--> on 401 + X-Admin-Session-Invalid: clear session and navigate /login
    |
    v
Controller
    |
    v
Service layer
    |
    v
MongoDB / storage / response
""")

    return doc


def vietnamese_doc():
    doc = Document()
    configure_document(doc, "Tài liệu kỹ thuật dự án UIWEB")
    add_title_page(
        doc,
        "Tài liệu kỹ thuật dự án UIWEB",
        "Kiến trúc, workflow, lịch sử tính năng, code board, và định hướng tiếp theo",
        "Bản tiếng Việt",
        "AdminSite API, AdminSite Frontend, UserSite, SharedComponents, Contracts, Resource Manager, Content Management, Page Builder, tools, và các workflow vận hành.",
    )

    add_h1(doc, "1. Tổng quan điều hành")
    add_para(doc, "UIWEB là một hệ thống CMS và website gồm ba ứng dụng chính. AdminSite là dashboard để admin đăng nhập, xây trang, quản lý content, quản lý Resource Library, cấu hình branding/theme, quản lý user, form, và dữ liệu hoạt động. UserSite là website public. API là backend trung tâm phụ trách authentication, MongoDB, publish/reset, content workflow, resource management, public rendering data, form, metrics, và asset storage.")
    add_para(doc, "Ý tưởng kiến trúc quan trọng nhất là tách draft và published. Admin user chỉnh sửa draft collections. Khi publish, hệ thống clone page graph hoặc content item sang published collections. UserSite chỉ đọc published data thông qua public endpoints. Vì vậy public website không vô tình hiển thị bản nháp.")
    add_para(doc, "Ý tưởng quan trọng thứ hai là tách Direct Upload và Managed Resource. Direct Upload là asset dùng một lần, nằm trực tiếp trên field của page/section/block/content/branding. Managed Resource là asset tái sử dụng, nằm trong Resource Library và có thể được chọn ở nhiều nơi. Delete/replace asset phải đi qua usage check và cleanup để tránh xóa nhầm file còn đang dùng.")
    add_bullets(doc, [
        "AdminSite-API là nơi giữ business rule và persistence.",
        "AdminSite-Frontend là Blazor Server dashboard cho Editor, Manager, AdminAdmin.",
        "UserSite là Blazor Server public website dùng public API DTO.",
        "SharedComponents chứa renderer dùng chung cho public page và admin preview.",
        "Contracts chứa DTO dùng chung giữa API, AdminSite, UserSite, và SharedComponents.",
        "MongoDB lưu structured data và asset metadata. Binary file hiện đang lưu trong R2-compatible object storage, không lưu trực tiếp trong MongoDB.",
    ])
    add_code_board(doc, "System Board", """
[Admin User]
    |
    v
AdminSite-Frontend (Blazor Server dashboard)
    |
    v
AdminSite-API (ASP.NET Core Web API)
    |
    +--> MongoDB collections: draft, published, users, resources, forms, settings
    |
    +--> R2StorageService: uploaded images, files, videos
    |
    v
Public API endpoints under /api/public
    |
    v
UserSite + SharedComponents renderer
    |
    v
Public website visitors
""")

    add_h1(doc, "2. Bản đồ repository và module")
    add_table(doc, ["Folder", "Nhiệm vụ", "Ghi chú quan trọng"], [
        ("AdminSite-API", "ASP.NET Core backend. Chứa controllers, services, MongoDB collections, auth, publish/reset, content workflow, resource management, public API, forms, metrics, và asset storage.", "Project file: AdminSite-API/Main-API.csproj. Root namespace là FullProject."),
        ("AdminSite-Frontend", "Blazor Server admin dashboard. Chứa UI cho Page Builder, Content Management, Resource Library, global settings, users, forms, metrics.", "HttpService gọi API; AdminAuthService giữ admin_session trong localStorage."),
        ("UserSite", "Public Blazor Server website. Load navigation, theme, pages, content pages, footer, social, global buttons từ API.", "Dùng PublicApiService và SharedComponents/PageRenderer.razor."),
        ("SharedComponents", "Renderer và visual sections/blocks dùng chung.", "Quan trọng vì admin preview và UserSite nên render cùng data shape."),
        ("Contracts", "Shared DTO contracts và global helper DTOs.", "Contracts/Public điều khiển renderer DTO; Contracts/Auth định nghĩa roles và permissions."),
        ("Tool", "Actual project tools dùng bởi owner/tester.", "Có demo database import tool và CloneUtility coverage tool."),
        ("AI-Tools", "Helper scripts/reports chủ yếu dùng bởi Codex khi audit.", "Tách khỏi user-facing project tools để rõ nghĩa."),
        ("build-check", "Local build output để tránh DLL bị Visual Studio lock.", "Nên giữ local/ignored."),
    ], [1900, 4560, 2900])
    add_code_board(doc, "Runtime Application Board", """
AdminSite-Frontend
    Pages/*.razor, Services/*.cs
    -> calls API via HttpService
    -> stores admin_session in browser localStorage

AdminSite-API
    Controllers/*.cs
    -> Services/*.cs
    -> MongoDbContext collections
    -> R2StorageService for binary assets

UserSite
    Components/Pages/PageView.razor
    -> PublicApiService
    -> SharedComponents/PageRenderer.razor
    -> SharedComponents/Sections and Blocks
""")

    add_h1(doc, "3. Startup và cấu hình")
    add_para(doc, "API startup nằm trong AdminSite-API/Program.cs. File này đọc các section MongoDb, Jwt, Seed, Cors, R2Storage, và FormSecurity từ appsettings.json. Nếu MongoDB setting thiếu, app dừng ngay. Nếu JWT secret thiếu hoặc quá ngắn, app cũng dừng. Đây là bảo vệ cần thiết để backend không chạy với cấu hình auth sai.")
    add_para(doc, "Services được đăng ký theo feature: auth, pages, sections, blocks, publish/reset, content services, managed resource services, form services, public services, metrics, settings, asset cleanup, và R2 storage. MongoClient và IMongoDatabase là singleton, còn phần lớn application service là scoped.")
    add_para(doc, "Khi startup, API gọi MongoIndexService.EnsureIndexesAsync để tạo index. Nếu index creation lỗi, hệ thống log warning và tiếp tục chạy. Điều này giúp local testing không bị chặn hoàn toàn vì lỗi index tạm thời, nhưng query có thể chậm hơn cho tới khi index được xử lý.")
    add_para(doc, "Startup cũng xóa legacy default Resource Library albums và đảm bảo default public form definitions tồn tại. Admin seeding là optional, chỉ chạy nếu Seed:AdminEmail và Seed:AdminPassword được cấu hình.")
    add_code_board(doc, "API Startup Board", """
Program.cs
    |
    +--> read appsettings: MongoDb, Jwt, Cors, R2Storage, FormSecurity
    |
    +--> register MongoClient + IMongoDatabase + MongoDbContext
    |
    +--> register services by feature module
    |
    +--> configure JWT authentication + fallback authorization
    |
    +--> configure rate limits: public-form, admin-login
    |
    +--> EnsureIndexesAsync()
    |
    +--> RemoveLegacyDefaultAlbumsAsync()
    |
    +--> EnsureDefaultDefinitionsAsync()
    |
    +--> optional SeedAdminAsync()
    |
    v
Middleware: CORS -> HTTPS -> RateLimiter -> Authentication -> AdminSessionValidationMiddleware -> Authorization -> Controllers
""")

    add_h1(doc, "4. MongoDB data model")
    add_para(doc, "MongoDbContext là lớp expose các collection chính. Quyết định kiến trúc quan trọng nhất là tách draft/published. Pages, sections, blocks, và content items được chỉnh sửa trong draft collections, sau đó copy sang published collections khi publish. Public website chỉ đọc published collections.")
    add_table(doc, ["Nhóm collection", "Collections", "Mục đích"], [
        ("Page Builder", "pages_draft, pages_published, page_revisions", "Lưu page metadata, slug tree, SEO, page card, draft state, và published snapshots."),
        ("Sections và Blocks", "sections_draft, sections_published, blocks_draft, blocks_published", "Lưu layout sections và reusable block content."),
        ("Content", "content_draft, content_published, content_types, content_revisions, content_audit_logs", "Lưu articles, whitepapers, video resources, image resources, gallery resources, workflow history."),
        ("Resource Library", "managed_resources, resource_albums", "Lưu reusable asset metadata, album organization, source type, storage key, usage-protected records."),
        ("Admin Auth", "admin_users, admin_sessions, admin_login_activity, admin_audit_logs", "Lưu admin accounts, sessions, login records, audit records."),
        ("Forms", "form_definitions, form_submissions", "Lưu reusable form definitions và public submissions."),
        ("Global Site", "branding, site_settings, theme, footer, social, global_buttons", "Lưu cấu hình visual/navigation toàn site."),
        ("Metrics", "visitor_metrics", "Lưu page/content/download counters cho Website Activity."),
    ], [1750, 3400, 4210])
    add_para(doc, "MongoDB lưu metadata và content structures. Nó không lưu binary bytes của image/video/file trong application path thông thường. File upload hiện đi qua R2StorageService; MongoDB chỉ lưu URL, storage key, file name, MIME type, size, source, và resource identity.")

    add_h1(doc, "5. Authentication và Authorization")
    add_para(doc, "Authentication dùng JWT kết hợp server-side session validation. Một JWT hợp lệ chưa đủ: token phải khớp với active admin session trong MongoDB. Nhờ vậy hệ thống có thể revoke session khi logout, reset password, đổi role, đổi permission, disable user, hoặc delete account.")
    add_para(doc, "Roles nằm trong Contracts/Auth/AuthDtos.cs: AdminAdmin, Manager, Writer, Viewer. Permissions là page-builder, manage-content, publish-content, manage-users, manage-settings, delete-content, view-logs. AdminAdmin là full-access role và bypass permission check qua AdminAuthorization.HasPermission.")
    add_para(doc, "AuthService.LoginAsync normalize email, kiểm tra status/lockout, verify password bằng BCrypt, tạo token id, tạo JWT claims, lưu AdminSessionRecord, reset failed login counter, và ghi login/audit records. Login sai tăng FailedLoginAttempts và có thể lock account 15 phút sau 5 lần sai.")
    add_para(doc, "AdminSessionValidationMiddleware chạy sau JWT authentication. Nó bỏ qua public endpoints, Swagger, login, và health. Với request khác, nó đọc adminId, jti, tokenVersion từ JWT claims rồi gọi AuthService.ValidateSessionAsync. Session sai trả 401 kèm X-Admin-Session-Invalid: true. AdminSite HttpService thấy header này sẽ clear session và đưa user về /login.")
    add_code_board(doc, "Authentication Board", """
Login.razor
    |
    v
AuthController.Login
    |
    v
AuthService.LoginAsync
    |
    +--> admin_users: find email, status, lockout, password hash
    +--> admin_sessions: insert session with token id + token version
    +--> admin_login_activity/admin_audit_logs: record event
    |
    v
JWT returned to AdminSite
    |
    v
AdminAuthService stores admin_session in localStorage
    |
    v
HttpService attaches Bearer token to API requests
    |
    v
JWT middleware validates token signature/lifetime
    |
    v
AdminSessionValidationMiddleware validates MongoDB session
    |
    v
Controllers apply role/permission checks
""")
    add_table(doc, ["Role", "Khả năng mặc định", "Giới hạn quan trọng"], [
        ("AdminAdmin", "Full access: users, settings, page builder, content, publish/delete, logs, Resource Library.", "Không được xóa/disable AdminAdmin active cuối cùng."),
        ("Manager", "Quản lý content, publish content, delete content mặc định.", "User management và settings cần permission riêng."),
        ("Writer", "Manage content và được vào Resource Library theo hướng hiện tại.", "Không tự động publish hoặc manage users."),
        ("Viewer", "Read-only dashboard role.", "Không nên xem Resource Library theo direction hiện tại."),
    ], [1700, 4800, 2860])

    add_h1(doc, "6. Page Builder workflow")
    add_para(doc, "Page Builder xoay quanh Page, Section, và Block documents. Page định nghĩa navigation và identity của public page. Section định nghĩa cấu trúc dọc của page. Block định nghĩa các đơn vị content reusable bên trong section hoặc column. Draft records được edit trong AdminSite; published records chỉ được tạo bởi PublishService.")
    add_para(doc, "Các identity field quan trọng là StableId, SourceId, Version, PublishedAt. StableId là logical identity dài hạn giữa draft và published. SourceId cho biết draft document nào tạo ra published snapshot. Version tăng khi publish. Public rendering thường dựa vào PageStableId và SectionStableId, không chỉ MongoDB document id.")
    add_para(doc, "ColumnsSection cần chú ý đặc biệt. Column slots có IDs, và blocks trong columns trỏ tới slot qua ColumnSlotId. CloneUtility phải giữ ColumnSlot.Id và Block.ColumnSlotId khi publish/reset, nếu không block có thể render sai column hoặc biến mất.")
    add_code_board(doc, "Page Builder Edit and Publish Board", """
Canvas.razor / EditPanel.razor
    |
    +--> AdminPageService -> PagesController -> PageService -> pages_draft
    +--> AdminSectionService -> SectionsController -> SectionService -> sections_draft
    +--> AdminBlockService -> BlocksController -> BlockService -> blocks_draft
    |
    v
User clicks Publish
    |
    v
PagesController.POST /api/admin/pages/{pageId}/publish
    |
    v
PublishService.PublishPageAsync
    |
    +--> load draft page, sections, blocks
    +--> delete existing published page graph for same StableId
    +--> CloneUtility.ClonePage / CloneSection / CloneBlock
    +--> insert into pages_published, sections_published, blocks_published
    +--> mark draft page Published and increment Version
    +--> after commit: AssetCleanupService removes unused replaced assets
""")
    add_table(doc, ["File", "Nhiệm vụ", "Vì sao quan trọng"], [
        ("AdminSite-API/Controllers/PagesController.cs", "Endpoints cho CRUD page, publish, reset, hierarchy.", "Entry point cho phần lớn Page Builder commands."),
        ("AdminSite-API/Services/PagesServices.cs", "Page persistence, slug/full slug, retrieval, children, public lookup.", "Giữ page identity và navigation nhất quán."),
        ("AdminSite-API/Services/SectionServices/SectionServices.cs", "Section creation, update, ordering, visibility, type-specific behavior.", "Sở hữu data shape của sections."),
        ("AdminSite-API/Services/BlockService.cs", "Block create/update/delete/order và public block retrieval.", "Sở hữu block data và column/container relationships."),
        ("AdminSite-API/Services/PublishAndResetService/PublishService.cs", "Draft-to-published transaction.", "Tạo public snapshot."),
        ("AdminSite-API/Services/PublishAndResetService/ResetService.cs", "Published-to-draft reset transaction.", "Khôi phục draft từ public snapshot."),
        ("AdminSite-API/Utils/CloneUtility.cs", "Manual deep clone của page/section/block graph.", "Design debt vì field mới có thể bị thiếu khi clone."),
    ], [3100, 3300, 2960])

    add_h1(doc, "7. CloneUtility và điểm yếu thiết kế hiện tại")
    add_para(doc, "CloneUtility hiện copy thủ công theo từng concrete type. Nó biết mọi section type và block type, sau đó tạo object mới và copy từng field. Cách này chạy được nếu bảo trì kỹ, nhưng dễ lỗi: nếu developer thêm field mới vào model mà quên thêm vào CloneUtility, publish/reset có thể âm thầm làm mất field đó.")
    add_para(doc, "Một lỗi/điểm yếu thực tế đã tìm ra là ColumnSlotId. ColumnsSection dùng slot IDs và block trỏ tới slot đó. Nếu clone không giữ những IDs này, published page có thể khác draft page. Cleanup gần đây đã tăng coverage và bảo vệ hành vi column slot, nhưng thiết kế tổng thể vẫn nên thay.")
    add_para(doc, "Hướng tương lai là PageGraphCloneService hoặc clone-profile system. Thay vì một switch statement thủ công phải biết mọi field mãi mãi, clone nên dùng serializer/profile approach với patch steps rõ ràng cho identity fields. Profiles có thể là PublishSnapshot, DraftResetSnapshot, PresetApply, DuplicateAsNewContent.")
    add_code_board(doc, "Future Clone Design Board", """
Current:
    CloneUtility.CloneSection(section)
        switch on concrete type
        manually copy every field
        risk: new field can be forgotten

Better:
    PageGraphCloneService.Clone(graph, profile)
        deep-copy by serialization or mapping profile
        apply identity patch rules:
            - new MongoDB Id
            - preserve or create StableId based on profile
            - set SourceId
            - set PublishedAt
            - set Status
            - set Version
        run coverage/audit for fields excluded by profile
""")

    add_h1(doc, "8. Content Management workflow")
    add_para(doc, "Content Management tách khỏi Page Builder. Nó quản lý article-like pages và resource-like content records như whitepaper, video, image, gallery. Cải tiến lớn là thêm content behavior. Một content type có thể là Page, FileResource, VideoResource, ImageResource, hoặc Gallery. Behavior điều khiển editor fields, validation, preview, và LibrarySection rendering.")
    add_para(doc, "ContentService hiện đóng vai facade trên các service nhỏ hơn. User đã quyết định không split nửa vời quá sớm, nhưng các module quan trọng đã có: ContentTypeService lo type CRUD/defaults; ContentValidationService lo behavior-aware validation; ContentWorkflowService lo submit/publish/delete/restore/permanent delete; ContentRevisionService lo revisions/audit logs; ContentMappingService lo mapping/published clone; ContentAssetMetadataService normalize ResourceId, ResourceSource, StorageKey, body media, gallery items, attachments.")
    add_para(doc, "Page behavior cần body content. FileResource cần file. VideoResource cần video. ImageResource cần image. Gallery cần gallery items. Rule hiện tại: Content Page Preview chỉ render Page behavior. Resource-based content được theo dõi qua All Content/My Content và Resource/Library flows, không giả lập thành article page.")
    add_code_board(doc, "Content Create/Update Board", """
ContentEditorShell.razor / ContentItems.razor
    |
    v
AdminContentService
    |
    v
ContentController
    |
    v
ContentService
    |
    +--> ContentValidationService: behavior-aware rules
    +--> ContentAssetMetadataService: ResourceId/ResourceSource/StorageKey normalization
    +--> ContentRevisionService: save revision before update
    +--> MongoDB content_draft
    +--> AssetCleanupService: delete replaced direct-upload assets if unused
""")
    add_code_board(doc, "Content Publish/Delete Board", """
ContentController.Publish/Delete/Restore/PermanentDelete
    |
    v
ContentWorkflowService
    |
    +--> Start MongoDB transaction
    +--> Validate publish rules
    +--> Update content_draft status
    +--> CloneForPublished via ContentMappingService
    +--> Replace content_published record by StableId
    +--> Write revision/audit log
    +--> Commit transaction
    |
    v
AssetCleanupService deletes unused replaced assets after commit
""")
    add_table(doc, ["Behavior", "Rule trong editor", "Rule preview/display"], [
        ("Page", "Title, summary, cover/hero, body blocks/html được bật.", "Content Page Preview và public content detail page."),
        ("FileResource", "Managed file selector/upload bắt buộc; body ẩn.", "LibrarySection mở file ở tab mới; All/My Content mở file trực tiếp."),
        ("VideoResource", "Video upload hoặc YouTube video link hợp lệ bắt buộc; body ẩn.", "LibrarySection mở video modal/player."),
        ("ImageResource", "Image selector/upload bắt buộc; body ẩn.", "LibrarySection mở image lightbox/modal."),
        ("Gallery", "Gallery items sẽ bắt buộc; body ẩn.", "LibrarySection Gallery layout hiển thị gallery."),
    ], [1750, 4010, 3600])

    add_h1(doc, "9. Resource Manager và Resource Library")
    add_para(doc, "Resource Manager đã được overhaul để trở thành một library thật sự, không phải một màn preview content thứ hai. Nó quản lý reusable assets: images, files, videos. Content Management chọn những file/media để public display, còn Resource Library quản lý pool lớn hơn. Hai chức năng liên quan nhưng không trùng nhau.")
    add_para(doc, "Rule hiện tại là một resource chỉ nằm trong tối đa một album. Album chỉ là organizational metadata. Album không giới hạn nơi resource được reuse. Một resource có thể dùng trên nhiều pages, sections, blocks, branding fields, và content records dù chỉ thuộc một album.")
    add_para(doc, "Album tách scope. Media albums chứa images và videos. File albums chứa files. Nếu cần nhóm phục vụ một page, dùng naming convention như File-Page1 và Resource-Page1 thay vì mixed album. Delete album là hard-block: nếu album còn resources, user phải move hoặc unfile resources trước khi xóa album.")
    add_para(doc, "Direction UI của Resource Library là main list/grid gallery ở trung tâm và compact sidebar. Main area dùng để browse/select resources. Sidebar dùng cho details: preview/file icon, name, description, Preview action, usage, file size, MIME type, uploaded by, created time, updated time. Internal fields như storage key và raw URL không nên hiện cho non-code users.")
    add_code_board(doc, "Resource Library Board", """
ContentResources.razor
    |
    +--> Normal mode: resource list/grid, type tabs, search, view switch
    +--> Album mode: folder-like albums, scoped to media or file
    +--> Add Resource: type-first modal, upload queue, optional YouTube link for video
    +--> Add Album/Group: create organizational folder
    |
    v
AdminContentService
    |
    v
ManagedResourcesController
    |
    +--> ManagedResourceService
        +--> ManagedResourceValidationService
        +--> ManagedResourceAlbumService
        +--> ManagedResourceUsageService
        +--> AssetCleanupService
""")
    add_code_board(doc, "Resource Delete Board", """
User clicks Delete Resource
    |
    v
ManagedResourcesController.Delete
    |
    v
ManagedResourceService.DeleteAsync
    |
    +--> ManagedResourceUsageService.GetUsageAsync(resource)
    |
    +--> if UsageCount > 0:
    |       return error: remove references before deleting
    |
    +--> if UsageCount == 0:
            delete managed_resources record
            AssetCleanupService.DeleteUnusedAsync(resource.Url, resource.ThumbnailUrl)
            AssetReferenceService verifies no remaining references
            R2StorageService deletes storage object
""")
    add_table(doc, ["Resource Type", "Storage rule", "UI rule"], [
        ("Image", "Binary upload đi vào R2/local storage abstraction. Mongo lưu URL, StorageKey, file metadata, album id, usage metadata.", "Hiển thị list/grid image. Có thể upload mới hoặc chọn từ Resource Library."),
        ("File", "Binary upload đi vào storage. Mongo lưu URL, StorageKey, filename, MIME, size.", "Grid dùng file icon/placeholder theo extension. Không cần editable preview image."),
        ("Video Upload", "Video file upload đi vào storage. Mongo lưu URL, StorageKey, filename, MIME, size. Accepted formats gồm mp4, webm, mov.", "Không custom thumbnail. Dùng video label/name để navigation."),
        ("Video YouTube Link", "Ngoại lệ chỉ cho video. Không upload R2. Mongo lưu normalized YouTube embed URL và Source = external-url.", "Chỉ nhận real YouTube video links, không nhận playlist/channel/search pages."),
    ], [1800, 4700, 2860])

    add_h1(doc, "10. Resource Picker và asset metadata")
    add_para(doc, "Resource Picker phải lưu metadata, không chỉ copy URL. Các field quan trọng là ResourceId, ResourceSource, StorageKey, cùng với visible URL. Nhờ vậy hệ thống biết asset là DirectUpload hay ManagedResource, track usage, propagate replacement, và cleanup old assets an toàn.")
    add_para(doc, "DirectUpload nghĩa là asset gắn vào một field cụ thể và không hiện trong Resource Library. ManagedResource nghĩa là asset reusable trong library. Hero backgrounds và decorative section/block images thường default DirectUpload. Whitepapers, reports, tool files, gallery items, reusable video/image resources default ManagedResource.")
    add_code_board(doc, "Picker Metadata Board", """
ResourcePickerModal.razor
    |
    v
User selects resource
    |
    v
Target editor field receives:
    Url
    ResourceId
    ResourceSource = ManagedResource
    StorageKey
    FileName / Mime / Size when relevant
    |
    v
ContentAssetMetadataService or section/block update mapper normalizes values
    |
    v
MongoDB stores both visible URL and source metadata
""")

    add_h1(doc, "11. Asset storage và cleanup")
    add_para(doc, "Ban đầu dự án từng nghĩ tới base64-style asset storage, nhưng hướng đó làm database quá nặng. Thiết kế hiện tại lưu binary assets bên ngoài MongoDB qua R2StorageService. MongoDB lưu references và metadata. Đây cũng là hướng đúng cho company-local storage sau này: thay implementation storage nhưng giữ metadata contract.")
    add_para(doc, "AssetCleanupService là rule cleanup trung tâm. Nó nhận old URLs từ records bị replace/delete, hỏi AssetReferenceService xem URL còn được reference ở đâu không, rồi mới gọi R2StorageService.DeleteAsync. Cách này bảo vệ shared resources và tránh xóa file còn đang dùng.")
    add_para(doc, "Resource Library delete dùng cùng global rule. Nếu managed resource còn được dùng, không được delete. Nếu unused và user chọn erase, record bị xóa và storage object được cleanup. Direct-upload assets được cleanup khi replace hoặc khi owning page/content graph bị delete, miễn là không còn record nào reference cùng URL.")
    add_code_board(doc, "Asset Cleanup Board", """
Old asset URL detected
    |
    v
AssetCleanupService.DeleteIfUnusedAsync(oldUrl, replacementUrl)
    |
    +--> skip if oldUrl is empty
    +--> skip if oldUrl == replacementUrl
    +--> AssetReferenceService.IsReferencedAsync(oldUrl)
    |
    +--> if referenced:
    |       keep storage object
    |
    +--> if not referenced:
            R2StorageService.DeleteAsync(oldUrl)
""")
    add_table(doc, ["Storage Layer", "Hiện tại", "Hướng company-local"], [
        ("Binary storage", "Cloudflare R2/S3-compatible storage qua R2StorageService.", "Implement local storage service tương đương trả về URL/path và storage key."),
        ("Database", "MongoDB lưu metadata, không lưu bytes: URL, StorageKey, ResourceId, ResourceSource, MIME, size.", "Giữ metadata model nếu có thể để tránh rewrite editors/renderers."),
        ("Cleanup", "AssetCleanupService và AssetReferenceService bảo vệ asset còn dùng.", "Reuse cùng cleanup service với company-local storage."),
        ("Security", "Đã có extension, MIME, folder, signature checks. Anti-virus/deep scan là future work.", "Thêm scan/quarantine pipeline trước khi lưu final."),
    ], [1800, 3700, 3860])

    add_h1(doc, "12. Public rendering và UserSite")
    add_para(doc, "Public website không đọc admin draft collections trực tiếp. UserSite gọi /api/public endpoints. PublicController giao page assembly cho PublicPageAssemblyService và metadata cho PublicMetadataService. Response được map thành shared public DTOs, rồi render bằng SharedComponents/PageRenderer.razor.")
    add_para(doc, "PublicPageAssemblyService lấy published page, visible published sections, và visible published blocks. Nó xử lý special sections trước. ShowcaseSection đọc published child pages và map thành cards. LibrarySection đọc published content items theo content types được chọn và map thành resource-aware library items. ColumnsSection group blocks theo ColumnSlotId. Standard sections đi qua catch-all branch.")
    add_para(doc, "UserSite/PublicApiService nhận JSON, map section types thành Contracts.Public DTO classes, build block tree, map column slot blocks, và đưa PublicPageDto cuối cùng vào PageRenderer. PageRenderer pattern match từng PublicSectionDto subtype và gọi component tương ứng trong SharedComponents/Sections.")
    add_code_board(doc, "Public Page Render Board", """
Visitor navigates to /solutions
    |
    v
UserSite PageView.razor
    |
    v
PublicApiService.GetPageAsync("solutions")
    |
    v
GET AdminSite-API /api/public/pages/solutions
    |
    v
PublicController.GetPage
    |
    v
PublicPageAssemblyService.GetPageResponseAsync
    |
    +--> PageService.GetByFullSlugAsync
    +--> SectionService.GetPublicSectionsByPageAsync
    +--> BlockService.GetPublicByPageAsync
    +--> map special sections and standard sections
    |
    v
Contracts.Public DTOs
    |
    v
SharedComponents/PageRenderer.razor
    |
    v
HeroSection, LibrarySection, ColumnsSection, etc.
""")

    add_h1(doc, "13. LibrarySection rendering")
    add_para(doc, "LibrarySection trở thành renderer chính cho managed content resources. Nó hỗ trợ layouts hiện có: Card, Grid, Rows, List, Gallery. Resource behavior quyết định CTA làm gì. FileResource mở file ở tab mới. VideoResource mở video modal/player. ImageResource mở lightbox/modal. Gallery hiển thị gallery layout và mở images/videos bằng modal/lightbox.")
    add_para(doc, "Điều này thay thế overlap cũ với GallerySection. Direction là thêm Library Gallery layout trước, giữ old renderer cho old pages nếu cần, hide GallerySection khỏi Add Section, rồi cleanup sau migration. Resource-based content không nên giả làm article page bình thường.")
    add_code_board(doc, "LibrarySection Behavior Board", """
LibrarySectionEditor.razor
    |
    v
LibrarySection stored in sections_draft/published
    |
    v
PublicPageAssemblyService detects LibrarySection
    |
    v
ContentService.GetPublishedLibraryItemsAsync
    |
    v
MapLibraryItem(content, contentType)
    |
    v
SharedComponents/Sections/LibrarySection.razor
    |
    +--> Page content: link to content detail page
    +--> FileResource: open file URL in new tab
    +--> VideoResource: modal video player
    +--> ImageResource: lightbox/modal
    +--> Gallery: gallery display with modal/lightbox
""")

    add_h1(doc, "14. Forms và public submission security")
    add_para(doc, "Forms được lưu dưới FormDefinition records và có thể dùng bởi FormBlock hoặc modal forms. Startup gọi FormDefinitionService.EnsureDefaultDefinitionsAsync, nên common public modal forms có sẵn mà không cần setup thủ công. Public submissions đi qua public endpoints và được bảo vệ bởi rate limiting và request size limits.")
    add_para(doc, "PublicController expose form submit endpoints theo page/section/block path và modal form path. Các endpoint này AllowAnonymous vì public visitor phải submit được, nhưng dùng public-form rate limiting policy. FormSubmissionSecurityService và FormValidationService xử lý validation và abuse checks như honeypot/payload limits.")
    add_code_board(doc, "Public Form Board", """
Visitor submits form
    |
    v
SharedComponents/Blocks/FormBlock.razor
    |
    v
UserSite PublicApiService.SubmitFormAsync
    |
    v
POST /api/public/pages/{slug}/sections/{sectionId}/blocks/{blockId}/form/submit
    |
    v
PublicFormSubmissionHandler
    |
    +--> FormDefinitionService
    +--> FormValidationService
    +--> FormSubmissionSecurityService
    +--> FormSubmissionService
    |
    v
MongoDB form_submissions
""")

    add_h1(doc, "15. Theme, Branding, Navigation, và Global Modules")
    add_para(doc, "Global modules cấu hình những phần của public site không thuộc một page body cụ thể: branding, theme, navigation, footer, social buttons, global buttons. PublicMetadataService expose chúng qua /api/public endpoints. UserSite load chúng qua PublicApiService và dùng trong layout components như NavMenu, MainLayout, SiteHeader, SiteFooter, SocialFloat.")
    add_para(doc, "Direction gần đây cho Theme và UserSite layout là top page navigation bar nên lấy màu từ Theme thay vì hard-coded static color. Sections cũng có hoặc nên giữ Background Type tên Theme để section tự dùng theme color. Nav item alignment cũng cần/đã được chỉnh để nằm giữa visual bar.")
    add_code_board(doc, "Global Metadata Board", """
Admin Global UI
    |
    +--> Branding.razor -> BrandingController -> BrandingService
    +--> Theme.razor -> ThemeController -> ThemeService
    +--> Footer.razor -> FooterController -> FooterService
    +--> Social.razor -> SocialController -> SocialButtonsService
    +--> GlobalButtons.razor -> GlobalButtonsController -> GlobalButtonsService
    |
    v
MongoDB global collections/settings
    |
    v
PublicMetadataService
    |
    v
/api/public/branding, /theme, /footer, /social, /global-buttons, /navigation
    |
    v
UserSite layout + SharedComponents
""")

    add_h1(doc, "16. Upload security")
    add_para(doc, "UploadSecurityPolicy hiện kiểm tra allowed upload folders, file extensions, MIME types, và file signatures. Với managed resources, nó kiểm tra selected kind có khớp allowed formats không. Video signatures gồm mp4/mov ftyp và webm signature. Office files, PDF, images, và text files cũng được signature check.")
    add_para(doc, "Đây là baseline protection hữu ích, nhưng chưa phải anti-virus scanning. Future plan là thêm anti-virus/deep scanning và database-bomb prevention. Thiết kế an toàn hơn là upload quarantine pipeline: nhận file, validate extension/MIME/signature/size, scan file, chỉ lưu chính thức sau khi scan pass, rồi mới create/update database record.")
    add_code_board(doc, "Future Upload Scan Board", """
Upload request
    |
    v
Controller request size + selected kind
    |
    v
UploadSecurityPolicy: folder + extension + MIME + signature
    |
    v
Quarantine temporary storage
    |
    v
Anti-virus / deep scan / archive-bomb checks
    |
    +--> fail: delete quarantine file, return error
    |
    +--> pass: move to permanent storage, save Mongo metadata
""")

    add_h1(doc, "17. Metrics và Website Activity")
    add_para(doc, "VisitorMetricService tăng counters cho public page views, public content page views, và downloads. PublicController tăng metrics sau khi lookup page/content thành công, và expose endpoint nhỏ để track download. AdminSite page WebsiteActivity.razor hiển thị metrics cho admin xem.")
    add_para(doc, "Metrics hiện cố ý đơn giản. Nó hữu ích để thấy page/resource nào đang active mà không cần xây analytics platform lớn. Nếu sau này cần analytics sâu hơn, counter hiện tại vẫn có thể giữ như internal signal nhẹ, còn tool analytics chuyên dụng xử lý hành vi visitor chi tiết hơn.")

    add_h1(doc, "18. Demo Database Import Tool")
    add_para(doc, "Project có one-click import tool trong Tool/Tool For Demo Import. Tool này cố ý import-only. Nó import prepared JSON collections vào MongoDB database để local tester mới có thể reproduce current UserSite layout và admin demo data mà không cần database cá nhân của owner.")
    add_para(doc, "Tool có thể tạo target database khi import. Target name đã thảo luận là FullProjectDb-UIWEB-3. Tester không cần tạo database này trước. Họ cần chỉnh application appsettings.json để trỏ vào database đó khi muốn test bằng demo data vừa import.")
    add_para(doc, "Import tool không duplicate R2 files. Nó import metadata và URLs. Nếu sau này migration sang company-local storage, nên tạo project migration/import riêng sau khi đã biết chi tiết storage của công ty.")
    add_code_board(doc, "Demo Import Board", """
run-import.bat
    |
    v
DemoDbImporter.exe / DemoDbImporter project
    |
    +--> appsettings.importer.json: MongoDB connection + target database
    +--> demo-seed/manifest.json: collection order
    +--> demo-seed/collections/*.json: data snapshots
    |
    v
MongoDB target database, for example FullProjectDb-UIWEB-3
    |
    v
AdminSite-API appsettings.json must point to that database for testing
""")

    add_h1(doc, "19. Lịch sử dự án và direction")
    add_para(doc, "Dự án đã đi qua nhiều quyết định lớn. Ý tưởng asset ban đầu từng nghiêng về base64-like data, nhưng điều đó làm database rất nặng. Hệ thống chuyển sang external storage với URL/storage key metadata. Sau đó Resource Management được thêm để không phải upload nào cũng thành reusable resource. Nhờ vậy có hai path rõ: DirectUpload cho asset dùng một lần và ManagedResource cho reusable library assets.")
    add_para(doc, "LibraryFeature plan ban đầu có chín phase: define asset paths, content type behavior, Resource Management base, Resource Picker, Content Editor integration, LibrarySection integration, Media/Resource Preview, GallerySection deprecation, safety cleanup. Phase 10 sau đó là service split/cleanup, gồm ContentService decomposition. User đã nhấn mạnh muốn real service split hơn là partial cleanup nửa vời.")
    add_para(doc, "Resource Manager Overhaul 2 sau đó biến Resource Library thành UI library chuyên nghiệp hơn. Main list/grid là trung tâm. Sidebar là details và usage. Albums là folder-like organization, one resource per album, không auto Unsorted album, hard-block delete nếu album còn resource, và type-first upload modal với upload queue dễ hiểu hơn.")
    add_para(doc, "Direction hiện tại là stability và clarity. Các theme lớn: theo phase order, tránh overlapping screens, ẩn code/internal fields khỏi non-code admin, bảo vệ asset khỏi accidental deletion, cleanup unused storage objects, giữ Resource Library tập trung vào reusable assets, và cuối cùng redesign CloneUtility để publish không âm thầm thiếu field mới.")
    add_table(doc, ["Giai đoạn", "Thay đổi", "Ý nghĩa"], [
        ("Early asset direction", "Base64-style storage bị bỏ, chuyển sang URL/storage key metadata.", "Tránh MongoDB bloat và giữ storage replaceable."),
        ("LibraryFeature", "DirectUpload vs ManagedResource, behavior-driven content, picker metadata, LibrarySection rendering.", "Tách rõ one-off uploads và reusable resources."),
        ("Resource Manager Overhaul 2", "Central gallery/list UI, compact sidebar, albums, bulk upload queue, settings, role visibility.", "Làm Resource Library dễ hiểu hơn cho non-code users."),
        ("Content service cleanup", "ContentService facade cộng type, validation, workflow, revision, mapping, asset metadata services.", "Giảm độ nặng của một service lớn mà không split giả."),
        ("Asset cleanup", "AssetCleanupService verify references trước khi storage deletion.", "Giảm storage bloat và bảo vệ reused assets."),
        ("Demo import tool", "Import-only tool cho local demo database setup.", "Giúp local/company testing dễ hơn."),
        ("Current stability work", "Mongo indexes, CloneUtility coverage, workflow transactions, namespace/csproj consistency.", "Làm codebase chắc hơn trước migration lớn."),
    ], [1850, 4010, 3500])

    add_h1(doc, "20. Điểm yếu còn lại và future work")
    add_bullets(doc, [
        "CloneUtility nên được redesign thành PageGraphCloneService hoặc clone profile system an toàn hơn.",
        "Log Overhaul vẫn planned nhưng cố ý chưa chạm trong recent stability phases. Direction là tách Login Log và Audit Trail, có retention/archive/export thay vì xóa casual.",
        "Migration sang company-local storage cần thêm thông tin về database/storage environment của công ty. Nên thay R2StorageService phía sau abstraction thay vì rewrite business modules.",
        "Anti-virus/deep scanning/database-bomb prevention nên thêm vào upload flow.",
        "Models.cs và admin DTO/model splitting nên để sau và làm như structural cleanup thật, không split vội nửa vời.",
        "Resource Library settings nên admin-governed và không expose internal storage fields cho normal users.",
        "Testing nên tăng quanh publish/reset, CloneUtility field coverage, Resource Library delete/usage, content behavior validation, asset cleanup.",
        "Tool folder ownership nên rõ: Tool cho actual project tools, AI-Tools cho assistant/audit helper scripts.",
    ])

    add_h1(doc, "21. Developer orientation: cách thay đổi hệ thống an toàn")
    add_para(doc, "Khi thêm field mới vào Page, Section, Block, ContentItem, ManagedResource, hoặc DTOs, đừng chỉ update UI. Phải kiểm tra toàn lifecycle: create/update DTOs, server model, validation, mapping, clone/publish/reset, public DTO mapping, renderer, resource replacement propagation, asset cleanup, database importer seed, và tests.")
    add_code_board(doc, "New Field Safety Checklist Board", """
New field added
    |
    +--> Model: AdminSite-API/Models.cs
    +--> DTO: AdminSite-API/DTOs.cs or Contracts/*
    +--> Admin UI model: AdminSite-Frontend/Models/AdminModels.cs
    +--> Editor UI and service mapper
    +--> Controller create/update endpoint
    +--> Validation service
    +--> Publish/reset clone or content mapping
    +--> PublicPageAssemblyService mapping if public
    +--> Contracts.Public DTO if rendered
    +--> SharedComponents renderer if visible
    +--> AssetCleanupService if field stores asset URL
    +--> ManagedResourceService replacement propagation if field can use Resource Library
    +--> Demo importer seed if demo data requires it
    +--> Focused test or coverage tool
""")

    add_h1(doc, "22. Appendix: Code boards quan trọng")
    add_code_board(doc, "All Content File Link Behavior", """
All Content / My Content list row click
    |
    +--> If content behavior == Page:
    |       open content page/detail route
    |
    +--> If content behavior == FileResource:
    |       open file URL directly in new tab
    |
    +--> If content behavior == VideoResource/ImageResource/Gallery:
            use resource-aware action rather than article-page redirect
""")
    add_code_board(doc, "Video Background Direction", """
SectionStylePanel.razor
    |
    +--> Background Type: Video
    +--> Upload New Video
    +--> Choose From Resource Library
    |
    v
Section.Style.BackgroundVideoUrl + metadata
    |
    v
SharedComponents/Helpers/StyleHelper.cs
    |
    v
SectionBackgroundMedia.razor renders muted autoplay loop video

YouTube/external URLs remain for Video Blocks or content video fields, not section background.
""")
    add_code_board(doc, "Background Image Fit Direction", """
Section background image
    |
    +--> Fit = Cover
    |       fills section, may crop image
    |
    +--> Fit = Contain
    |       shows full image, may leave empty space
    |
    +--> Position = center/top/bottom/left/right
            controls visible focal area
""")
    add_code_board(doc, "Resource Replacement Propagation", """
Resource Editor: Replace File
    |
    v
ManagedResourceService.ReplaceUploadAsync
    |
    +--> validate same kind
    +--> update ManagedResource URL/storage/file metadata
    +--> propagate to:
            content_draft/content_published
            pages_draft/pages_published
            sections_draft/sections_published
            blocks_draft/blocks_published
            branding
    +--> AssetCleanupService.DeleteIfUnusedAsync(oldUrl, newUrl)
""")
    add_code_board(doc, "Admin Request Lifecycle", """
AdminSite UI event
    |
    v
Admin service class
    |
    v
HttpService.SendAsync
    |
    +--> attach JWT from localStorage
    +--> call API
    +--> parse ApiResponse<T>
    +--> on 401 + X-Admin-Session-Invalid: clear session and navigate /login
    |
    v
Controller
    |
    v
Service layer
    |
    v
MongoDB / storage / response
""")

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    english = english_doc()
    vietnamese = vietnamese_doc()
    english.save(OUT_DIR / "UIWEB_Project_Technical_Guide_EN.docx")
    vietnamese.save(OUT_DIR / "UIWEB_Project_Technical_Guide_VI.docx")
    print(OUT_DIR / "UIWEB_Project_Technical_Guide_EN.docx")
    print(OUT_DIR / "UIWEB_Project_Technical_Guide_VI.docx")


if __name__ == "__main__":
    main()
